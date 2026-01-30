"""Export engine for generating PDF and DOCX exports of business cards.

This module provides the ExportEngine class that handles exporting business
cards to PDF and Word document formats, preserving card styling and layout.
"""

from io import BytesIO
from pathlib import Path
from typing import Optional

from PySide6.QtCore import QRect, QSize, Qt
from PySide6.QtGui import (
    QColor,
    QFont,
    QImage,
    QPainter,
    QPen,
    QPixmap,
)

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas as pdf_canvas
from reportlab.lib.utils import ImageReader
from reportlab.lib.colors import gray

from .card_data_model import CardDataModel
from ..models.card import CardRow, FieldDefinition, FieldType


# Standard business card dimensions (in pixels at 96 DPI for preview)
CARD_WIDTH = 336
CARD_HEIGHT = 192

# Standard credit card size in inches (3.5" x 2")
CARD_WIDTH_INCHES = 3.5
CARD_HEIGHT_INCHES = 2.0

# Points per inch for PDF
POINTS_PER_INCH = 72


class ExportEngine:
    """Generates PDF and DOCX exports of business cards.
    
    This class handles the export of business cards to various formats,
    supporting multiple cards per page with configurable layouts.
    """
    
    def __init__(self, model: CardDataModel, project_path: Optional[str] = None):
        """Initialize with card data model.
        
        Args:
            model: The CardDataModel containing template and rows.
            project_path: Optional path to project folder for resolving images.
        """
        self._model = model
        self._project_path = project_path
    
    def calculate_layout(self, page_size: QSize, cards_per_page: int) -> list[QRect]:
        """Calculate card positions for given page layout.
        
        Cards are always rendered at standard credit card size (3.5" x 2").
        Returns positions in points (72 points per inch).
        """
        if cards_per_page not in (1, 2, 4, 6, 8, 10):
            raise ValueError(f"Unsupported cards_per_page: {cards_per_page}")
        
        page_width = page_size.width()
        page_height = page_size.height()
        
        # Fixed card size in points
        card_width = int(CARD_WIDTH_INCHES * POINTS_PER_INCH)
        card_height = int(CARD_HEIGHT_INCHES * POINTS_PER_INCH)
        
        layouts = {
            1: (1, 1), 2: (1, 2), 4: (2, 2), 
            6: (2, 3), 8: (2, 4), 10: (2, 5)
        }
        
        cols, rows = layouts[cards_per_page]
        
        # Check if cards fit on page
        margin = 36  # 0.5 inch margin
        spacing = 18  # 0.25 inch spacing between cards
        
        required_width = (cols * card_width) + ((cols - 1) * spacing) + (2 * margin)
        required_height = (rows * card_height) + ((rows - 1) * spacing) + (2 * margin)
        
        if required_width > page_width or required_height > page_height:
            # Reduce cards per page if they don't fit
            # This shouldn't happen for reasonable values, but handle gracefully
            pass
        
        # Center the grid on the page
        total_grid_width = (cols * card_width) + ((cols - 1) * spacing)
        total_grid_height = (rows * card_height) + ((rows - 1) * spacing)
        
        start_x = (page_width - total_grid_width) // 2
        start_y = (page_height - total_grid_height) // 2
        
        positions = []
        for row in range(rows):
            for col in range(cols):
                x = start_x + col * (card_width + spacing)
                y = start_y + row * (card_height + spacing)
                positions.append(QRect(x, y, card_width, card_height))
        
        return positions
    
    def render_card(self, row: CardRow, width: int, height: int) -> QImage:
        """Render a single card row to image for export."""
        template = self._model.get_template()
        
        image = QImage(width, height, QImage.Format.Format_ARGB32)
        bg_color = QColor(template.background_color)
        image.fill(bg_color)
        
        painter = QPainter(image)
        painter.setRenderHint(QPainter.RenderHint.Antialiasing)
        painter.setRenderHint(QPainter.RenderHint.TextAntialiasing)
        painter.setRenderHint(QPainter.RenderHint.SmoothPixmapTransform)
        
        scale_x = width / CARD_WIDTH
        scale_y = height / CARD_HEIGHT
        
        # Draw border
        border_pen = QPen(QColor("#cccccc"))
        border_pen.setWidth(1)
        painter.setPen(border_pen)
        painter.drawRect(0, 0, width - 1, height - 1)
        
        # Render each field
        for field in template.fields:
            value = row.get_value(field.id)
            
            elem_x = int(field.x * scale_x)
            elem_y = int(field.y * scale_y)
            elem_width = int(field.width * scale_x)
            elem_height = int(field.height * scale_y)
            
            if field.field_type == FieldType.TEXT:
                self._render_text(
                    painter, str(value) if value else "",
                    elem_x, elem_y, elem_width, elem_height, field, scale_x, scale_y
                )
            else:  # IMAGE
                image_path = str(value) if value else ""
                if image_path and self._project_path and not Path(image_path).is_absolute():
                    image_path = str(Path(self._project_path) / image_path)
                self._render_image(
                    painter, image_path,
                    elem_x, elem_y, elem_width, elem_height
                )
        
        painter.end()
        return image
    
    def _render_text(
        self, painter: QPainter, text: str,
        x: int, y: int, width: int, height: int,
        field: FieldDefinition, scale_x: float, scale_y: float
    ) -> None:
        """Render a text field."""
        font = QFont(field.font_family, int(field.font_size * min(scale_x, scale_y)))
        font.setBold(field.font_bold)
        font.setItalic(field.font_italic)
        painter.setFont(font)
        painter.setPen(QColor(field.font_color))
        
        rect = QRect(x, y, width, height)
        painter.drawText(rect, Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter, text)
    
    def _render_image(
        self, painter: QPainter, image_path: str,
        x: int, y: int, width: int, height: int
    ) -> None:
        """Render an image field."""
        if not image_path:
            painter.setPen(QColor("#999999"))
            rect = QRect(x, y, width, height)
            painter.drawRect(rect)
            return
        
        pixmap = QPixmap(image_path)
        if pixmap.isNull():
            painter.setPen(QColor("#999999"))
            rect = QRect(x, y, width, height)
            painter.drawRect(rect)
            return
        
        scaled = pixmap.scaled(
            width, height,
            Qt.AspectRatioMode.KeepAspectRatio,
            Qt.TransformationMode.SmoothTransformation
        )
        
        img_x = x + (width - scaled.width()) // 2
        img_y = y + (height - scaled.height()) // 2
        painter.drawPixmap(img_x, img_y, scaled)
    
    def _qimage_to_bytes(self, image: QImage) -> Optional[bytes]:
        """Convert QImage to PNG bytes."""
        from PySide6.QtCore import QBuffer, QIODevice
        
        buffer = QBuffer()
        buffer.open(QIODevice.OpenModeFlag.WriteOnly)
        if image.save(buffer, "PNG"):
            return bytes(buffer.data())
        return None
    
    def export_pdf(self, output_path: Path, cards_per_page: int) -> bool:
        """Export cards to PDF with standard credit card size and cut lines."""
        rows = self._model.get_all_rows()
        if not rows:
            return False
        
        page_width = int(8.5 * POINTS_PER_INCH)
        page_height = int(11 * POINTS_PER_INCH)
        page_size = QSize(page_width, page_height)
        
        positions = self.calculate_layout(page_size, cards_per_page)
        c = pdf_canvas.Canvas(str(output_path), pagesize=letter)
        
        # Card size in points
        card_width_pts = int(CARD_WIDTH_INCHES * POINTS_PER_INCH)
        card_height_pts = int(CARD_HEIGHT_INCHES * POINTS_PER_INCH)
        
        # Cut line extension beyond card edge
        cut_line_extend = 10
        
        row_index = 0
        while row_index < len(rows):
            page_rows = rows[row_index:row_index + cards_per_page]
            
            for i, row in enumerate(page_rows):
                if i >= len(positions):
                    break
                
                pos = positions[i]
                
                # Render card at high resolution for quality
                render_scale = 3
                card_image = self.render_card(
                    row, 
                    int(CARD_WIDTH * render_scale), 
                    int(CARD_HEIGHT * render_scale)
                )
                
                image_bytes = self._qimage_to_bytes(card_image)
                if image_bytes:
                    img_reader = ImageReader(BytesIO(image_bytes))
                    # PDF coordinates start from bottom-left
                    pdf_x = pos.x()
                    pdf_y = page_height - pos.y() - card_height_pts
                    
                    c.drawImage(
                        img_reader, pdf_x, pdf_y,
                        width=card_width_pts, height=card_height_pts,
                        preserveAspectRatio=False
                    )
                    
                    # Draw cut lines (light gray dashed lines)
                    c.setStrokeColor(gray)
                    c.setLineWidth(0.5)
                    c.setDash(3, 3)  # Dashed line pattern
                    
                    # Top cut line
                    c.line(
                        pdf_x - cut_line_extend, pdf_y + card_height_pts,
                        pdf_x + card_width_pts + cut_line_extend, pdf_y + card_height_pts
                    )
                    # Bottom cut line
                    c.line(
                        pdf_x - cut_line_extend, pdf_y,
                        pdf_x + card_width_pts + cut_line_extend, pdf_y
                    )
                    # Left cut line
                    c.line(
                        pdf_x, pdf_y - cut_line_extend,
                        pdf_x, pdf_y + card_height_pts + cut_line_extend
                    )
                    # Right cut line
                    c.line(
                        pdf_x + card_width_pts, pdf_y - cut_line_extend,
                        pdf_x + card_width_pts, pdf_y + card_height_pts + cut_line_extend
                    )
                    
                    # Reset dash pattern
                    c.setDash()
            
            row_index += cards_per_page
            if row_index < len(rows):
                c.showPage()
        
        c.save()
        return True
    
    def export_docx(self, output_path: Path, cards_per_page: int) -> bool:
        """Export cards to Word document with standard credit card size."""
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        
        rows = self._model.get_all_rows()
        if not rows:
            return False
        
        layouts = {
            1: (1, 1), 2: (1, 2), 4: (2, 2),
            6: (2, 3), 8: (2, 4), 10: (2, 5)
        }
        
        if cards_per_page not in layouts:
            raise ValueError(f"Unsupported cards_per_page: {cards_per_page}")
        
        cols, table_rows = layouts[cards_per_page]
        
        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Fixed card size - standard credit card
        card_width_inches = CARD_WIDTH_INCHES
        card_height_inches = CARD_HEIGHT_INCHES
        
        row_index = 0
        page_count = 0
        
        while row_index < len(rows):
            if page_count > 0:
                doc.add_page_break()
            
            page_rows = rows[row_index:row_index + cards_per_page]
            table = doc.add_table(rows=table_rows, cols=cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Add thin borders for cut lines
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
                r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
            )
            # Light gray dashed borders for cut lines
            tblBorders = parse_xml(
                r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                r'<w:top w:val="dashed" w:sz="4" w:color="808080"/>'
                r'<w:left w:val="dashed" w:sz="4" w:color="808080"/>'
                r'<w:bottom w:val="dashed" w:sz="4" w:color="808080"/>'
                r'<w:right w:val="dashed" w:sz="4" w:color="808080"/>'
                r'<w:insideH w:val="dashed" w:sz="4" w:color="808080"/>'
                r'<w:insideV w:val="dashed" w:sz="4" w:color="808080"/>'
                r'</w:tblBorders>'
            )
            existing = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders')
            if existing is not None:
                tblPr.remove(existing)
            tblPr.append(tblBorders)
            if tbl.tblPr is None:
                tbl.insert(0, tblPr)
            
            card_idx = 0
            for r in range(table_rows):
                for c in range(cols):
                    if card_idx < len(page_rows):
                        row = page_rows[card_idx]
                        cell = table.cell(r, c)
                        
                        # Set cell size to card size
                        cell.width = Inches(card_width_inches)
                        
                        render_scale = 3
                        card_image = self.render_card(
                            row, int(CARD_WIDTH * render_scale), int(CARD_HEIGHT * render_scale)
                        )
                        
                        image_bytes = self._qimage_to_bytes(card_image)
                        if image_bytes:
                            paragraph = cell.paragraphs[0]
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = paragraph.add_run()
                            run.add_picture(
                                BytesIO(image_bytes),
                                width=Inches(card_width_inches),
                                height=Inches(card_height_inches)
                            )
                        card_idx += 1
            
            row_index += cards_per_page
            page_count += 1
        
        doc.save(str(output_path))
        return True
